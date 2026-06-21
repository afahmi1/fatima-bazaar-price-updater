"""Invoice extraction.

Pulls candidate line items out of PDF, Excel (.xlsx/.xls/.csv) and Word (.docx)
invoices. Real-world invoices vary wildly, so this is intentionally heuristic:
it makes a best guess at (item name, quantity, unit/line cost) for each row and
the user confirms/edits everything on the review screen before prices are
computed. Nothing here is trusted blindly.

Each returned item is a dict:
    { "name": str, "qty": float|None, "cost": float|None, "raw": str }
where `cost` is the best-guess cost figure found on that line and `qty` is the
quantity if one was detected.
"""

import csv
import os
import re

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

MONEY_RE = re.compile(r"\$?\s*(\d{1,3}(?:,\d{3})*(?:\.\d+)?|\d+(?:\.\d+)?)")
# words that usually mean "this row is not a product line"
SKIP_WORDS = (
    "subtotal", "sub total", "total", "tax", "shipping", "freight", "invoice",
    "balance", "amount due", "thank you", "remit", "page ", "discount",
    "handling", "deposit", "credit", "account", "po number", "p.o.", "terms",
)


def _to_float(s):
    if s is None:
        return None
    s = str(s).replace(",", "").replace("$", "").strip()
    if s == "":
        return None
    try:
        return float(s)
    except ValueError:
        return None


def _looks_like_header_or_total(text: str) -> bool:
    low = (text or "").lower()
    return any(w in low for w in SKIP_WORDS)


def _money_values(text: str):
    """Return all money-ish numbers found in a string, as floats."""
    out = []
    for m in MONEY_RE.finditer(text or ""):
        v = _to_float(m.group(1))
        if v is not None:
            out.append(v)
    return out


# ---------------------------------------------------------------------------
# Column-aware row parsing (for tables: PDF tables, Excel rows, Word tables)
# ---------------------------------------------------------------------------

def _parse_table_rows(rows):
    """Given a list of rows (each a list of cell strings), try to detect the
    name / quantity / cost columns from a header, else fall back to heuristics.
    """
    items = []
    if not rows:
        return items

    # Try to find a header row to map columns.
    name_col = qty_col = cost_col = None
    header_idx = -1
    for i, row in enumerate(rows[:5]):
        cells = [str(c or "").strip().lower() for c in row]
        joined = " ".join(cells)
        if any(k in joined for k in ("description", "item", "product")):
            # Prefer the extended/line total column for cost (amount, ext, line
            # total) so that cost / units gives the true per-unit cost. Only
            # fall back to a unit-price/cost column if no total column exists.
            best_cost_rank = 99
            for j, c in enumerate(cells):
                if name_col is None and any(
                    k in c for k in ("description", "item", "product", "name")
                ):
                    name_col = j
                if qty_col is None and any(
                    k in c for k in ("qty", "quantity", "units", "cases", "ea")
                ):
                    qty_col = j
                rank = None
                if any(k in c for k in ("amount", "ext", "line total")):
                    rank = 0
                elif "total" in c:
                    rank = 1
                elif any(k in c for k in ("price", "cost")):
                    rank = 2
                if rank is not None and rank < best_cost_rank:
                    best_cost_rank = rank
                    cost_col = j
            header_idx = i
            break

    start = header_idx + 1 if header_idx >= 0 else 0
    for row in rows[start:]:
        cells = [str(c or "").strip() for c in row]
        if not any(cells):
            continue
        line_text = " ".join(c for c in cells if c)
        if _looks_like_header_or_total(line_text):
            continue

        name = None
        qty = None
        cost = None

        if name_col is not None and name_col < len(cells):
            name = cells[name_col].strip()
        if qty_col is not None and qty_col < len(cells):
            qty = _to_float(cells[qty_col])
        if cost_col is not None and cost_col < len(cells):
            cost = _to_float(cells[cost_col])

        # Fallbacks when no/partial header mapping.
        if not name:
            # longest mostly-alphabetic cell is probably the name
            text_cells = [c for c in cells if c and not _to_float(c)]
            name = max(text_cells, key=len) if text_cells else None
        if cost is None:
            nums = [_to_float(c) for c in cells if _to_float(c) is not None]
            # prefer the largest number on the row as the cost guess
            if nums:
                cost = max(nums)
        if qty is None:
            nums = [_to_float(c) for c in cells if _to_float(c) is not None]
            small = [n for n in nums if n and n == int(n) and n < 1000 and n != cost]
            if small:
                qty = small[0]

        if name and len(name) >= 2 and not name.replace(".", "").isdigit():
            items.append({
                "name": name,
                "qty": qty,
                "cost": cost,
                "raw": line_text,
            })
    return items


# Structured distributor-invoice rows like:
#   1.00  1.00  CS  10-119-2  RODOPA BULGARIAN FETA PL 12/400 GR  59.99  59.99
#   30.06 30.06 lb  10-195-2  M.E. NABULSI CHEESE 12/1 LB         6.99   210.14
# Columns: Qty Ordered, Qty Shipped, UOM, Item code, Description, Unit Price, Extended.
_ROW_RE = re.compile(
    r"^\s*(\d+(?:\.\d+)?)\s+(\d+(?:\.\d+)?)\s+([A-Za-z]{1,4})\s+"
    r"(\S+)\s+(.+?)\s+(\d{1,3}(?:,\d{3})*(?:\.\d+)?)\s+"
    r"(\d{1,3}(?:,\d{3})*(?:\.\d+)?)\s*$"
)
# Pack count embedded in a description: "12/400 GR" -> 12 units per case.
_PACK_RE = re.compile(r"(\d+)\s*/\s*[\d.]+\s*(?:GR|G|LB|OZ|KG|ML|L|CT|PC|EA)\b", re.I)
# Units of measure that mean the item is sold by weight.
_WEIGHT_UOMS = {"lb", "lbs", "kg", "oz", "g"}


def _parse_structured_invoice(text):
    """Parse regular distributor-invoice line rows (Qty/UOM/Item/Desc/Price/Ext).

    Returns [] if the text doesn't look like this format so callers can fall back.
    """
    items = []
    for line in (text or "").splitlines():
        m = _ROW_RE.match(line)
        if not m:
            continue
        qty_shp = _to_float(m.group(2))
        uom = m.group(3).lower()
        desc = m.group(5).strip()
        extended = _to_float(m.group(7))

        # Skip rows that weren't shipped / have no cost.
        if not extended or qty_shp in (None, 0):
            continue

        # Clean the name: drop asterisk noise ("FETA PL******", "**PKT**").
        name = re.sub(r"\*+", " ", desc)
        name = re.sub(r"\s{2,}", " ", name).strip()
        if len(name) < 2:
            continue

        by_weight = uom in _WEIGHT_UOMS
        if by_weight:
            # cost / pounds shipped = cost per lb
            divisor = qty_shp
        else:
            # cases shipped x units-per-case = total sellable units
            pack = _PACK_RE.search(desc)
            per_case = int(pack.group(1)) if pack else 1
            divisor = qty_shp * per_case

        items.append({
            "name": name,
            "qty": divisor,
            "cost": extended,
            "pricing_type": "weight" if by_weight else "quantity",
            "raw": line.strip(),
        })
    return items


# Ziyad-Brothers style cased-goods rows like:
#   4bx A0429A Al Afia Grape Leaves 12x32oz (908g) Jar 12ea/bx (≈7.83ea) 93.93 375.72
#   1bx 9791   Balparmak Flower Honey 12x16.2oz (460g) Jar 12ea/bx        87.14  87.14
# Columns: "<n>bx" cases, item#, description (with pack), Case$, Total.
_ZIYAD_RE = re.compile(r"^\s*(\d+)\s*bx\s+(\S+)\s+(.+)$", re.I)
_EA_PER_BX_RE = re.compile(r"(\d+)\s*ea\s*/\s*bx", re.I)
_MONEY_RE2 = re.compile(r"\d{1,3}(?:,\d{3})*\.\d{2}")
# Start of the pack/size portion of a description (used to trim the item name).
_PACK_START_RE = re.compile(
    r"\b\d+(?:\.\d+)?\s*(?:x|×)\s*\d"                       # 12x32, 6 x 850
    r"|\b\d+(?:\.\d+)?\s*(?:g|gm|gms|gram|grams|kg|oz|ml|l|lit|liter|litre|lb|fl)\b",
    re.I,
)


# Trailing packaging/container words to strip from item names ("...in glass",
# "...Packed in Bags", "...in tin"). Product descriptors like "frozen" are kept.
_PACKAGING_WORDS = {
    "in", "a", "glass", "tin", "tins", "jar", "jars", "box", "boxes", "pack",
    "packed", "packet", "packets", "foil", "bag", "bags", "bottle", "bottles",
    "can", "cans", "pouch", "pouches", "tub", "tubs", "tray", "trays",
    "carton", "cartons", "plastic", "pet", "container", "containers",
}


def _strip_packaging(name):
    """Drop trailing packaging words so 'Olive Oil in glass' -> 'Olive Oil'."""
    words = name.split()
    while words and words[-1].strip(",.;:").lower() in _PACKAGING_WORDS:
        words.pop()
    return " ".join(words).strip(" ,-")


def _parse_ziyad_invoice(text):
    """Parse cased-goods invoices priced per case (e.g. Ziyad Brothers).

    Per-item cost is figured from (cases x items-per-case): cost = line Total,
    divisor = cases * items_per_case, so cost/divisor is the true per-unit cost.
    """
    items = []
    for line in (text or "").splitlines():
        m = _ZIYAD_RE.match(line)
        if not m:
            continue
        cases = _to_float(m.group(1))
        rest = m.group(3)

        ea = _EA_PER_BX_RE.search(rest)
        if not ea:
            continue  # not a real cased-goods line
        per_case = int(ea.group(1))

        monies = _MONEY_RE2.findall(rest)
        if len(monies) < 2:
            continue
        total = _to_float(monies[-1])          # last number = line Total
        if not total:
            continue

        # Item name = description before the pack/size spec begins.
        cut = _PACK_START_RE.search(rest)
        name = (rest[:cut.start()] if cut else rest).strip()
        name = re.sub(r"\s{2,}", " ", name).strip(" ,-")
        name = _strip_packaging(name)
        if len(name) < 2:
            continue

        divisor = (cases or 1) * (per_case or 1)   # total individual units
        items.append({
            "name": name,
            "qty": divisor,
            "cost": total,
            "pricing_type": "quantity",
            "raw": line.strip(),
        })
    return items


# Karabetian-style rows:
#   0114 GOLDEN PLATE CYPRUS LEBNI 12/500g  1 738960197866 35.00 2.917 35.00
#   Item#  Description(+pack)  Qty  Barcode  Price(per case)  Unit(per item)  Amount
_KARA_MONEY = re.compile(r"\d{1,3}(?:,\d{3})*\.\d{2,3}")
_KARA_TAIL = re.compile(r"(\d+)(?:\s+(\d{8,14}))?\s*$")  # qty + optional barcode
_KARA_PACK = re.compile(
    r"\b\d+\s*/\s*\d"
    r"|\b\d+(?:\.\d+)?\s*(?:g|gr|gram|grams|kg|oz|ml|l|lit|lb|fl)\b",
    re.I,
)
# Pack ratio in a description: "24/100", "12/500g", "36x100x2 g". The FIRST
# number is the count of sellable units per case (e.g. 24 boxes); the rest is
# the size/contents of each unit.
_KARA_RATIO = re.compile(r"(\d+)\s*[/xX]\s*\d")
# Non-product charge lines to skip.
_KARA_SKIP = ("crv", "freight", "pallet", "delivery charge")


def _parse_karabetian(text):
    """Parse invoices with a per-item 'Unit' column and a Barcode column
    (e.g. Karabetian). Per-item cost = Amount / total units; the printed
    barcode is captured as the item's SKU."""
    lines = (text or "").splitlines()
    items = []
    for i, line in enumerate(lines):
        monies = list(_KARA_MONEY.finditer(line))
        if len(monies) < 3:
            continue
        # last three money figures = Price, Unit (per item), Amount
        unit = _to_float(monies[-2].group())
        amount = _to_float(monies[-1].group())
        if not amount or not unit or unit <= 0:
            continue
        prefix = line[: monies[-3].start()]  # text before Price

        tail = _KARA_TAIL.search(prefix.rstrip())
        if not tail:
            continue
        barcode = tail.group(2)
        head = prefix[: tail.start()].strip()
        if not head or any(w in head.lower() for w in _KARA_SKIP):
            continue

        # drop the leading item# code, then trim at the pack/size spec
        parts = head.split(None, 1)
        desc = parts[1] if len(parts) == 2 else head
        cut = _KARA_PACK.search(desc)
        name = (desc[: cut.start()] if cut else desc).strip()
        name = _strip_packaging(re.sub(r"\s{2,}", " ", name).strip(" ,-*"))
        if len(name) < 2:
            continue

        # barcode is sometimes wrapped onto the next line (digits only)
        if not barcode:
            for nxt in lines[i + 1 : i + 3]:
                m = re.match(r"^\s*(\d{8,14})\s*$", nxt)
                if m:
                    barcode = m.group(1)
                    break

        # Units per case from the description ratio: the FIRST number of "x/y"
        # is the count of sellable units (e.g. 24/100 TEA BAG = 24 boxes), not
        # the contents. Multiply by Qty (cases). This is more reliable than the
        # vendor's per-item column, which sometimes treats a whole case as one
        # unit. Fall back to Amount / Unit only when there's no ratio.
        qty = _to_float(tail.group(1)) or 1
        ratio = _KARA_RATIO.search(desc)
        if ratio:
            divisor = qty * int(ratio.group(1))
        else:
            divisor = round(amount / unit) or 1
        items.append({
            "name": name,
            "qty": divisor,
            "cost": amount,
            "pricing_type": "quantity",
            "barcode": barcode or "",
            "raw": line.strip(),
        })
    return items


# Nature's-Sweets style: a main line then a pack line on the NEXT line.
#   1 California Medjool Dates (Fancy)  6.00  70.00  420.00
#   2Lbx9 box                          -> 9 units/box, Qty=6 boxes -> div by 54
#   11 Healthy Trail mix              15.00  7.50  112.50
#   16Oz(452g) pcs                     -> sold as pieces, Qty IS the unit count
_NS_RE = re.compile(
    r"^(\d+)\s+(.+?)\s+([\d,]+\.\d{2})\s+([\d,]+\.\d{2})\s+([\d,]+\.\d{2})\s*$"
)
_NS_SIZE_PART = re.compile(r"\d+(?:\.\d+)?\s*(?:lbs?|oz|gr?|kg|ml|l)\b", re.I)


def _parse_pack_spec(spec):
    """Return (sold_as_pieces, units_per_box) from a pack line.

    The part of an "AxB" spec WITHOUT a weight unit is the count (units per
    box); the part with a weight unit is the size. "3Kg box" -> 1 unit/box.
    """
    low = spec.lower()
    is_pcs = bool(re.search(r"\bpcs?\b", low))
    core = re.sub(r"\b(box|pcs?|soleil)\b", "", spec, flags=re.I).strip()
    count = 1
    if re.search(r"[xX]", core):
        prod, found = 1, False
        for p in re.split(r"[xX]", core):
            p = p.strip()
            if not p or _NS_SIZE_PART.search(p):
                continue  # blank or a size (e.g. "2Lb", "110g")
            m = re.match(r"(\d+)", p)
            if m:
                prod *= int(m.group(1))
                found = True
        count = prod if found else 1
    return is_pcs, count


def _parse_natures_sweets(text):
    """Two-line items where Qty is boxes and the next line gives units-per-box.
    Per-item cost = Amount / (boxes x units-per-box), or / Qty for piece items.
    """
    lines = (text or "").splitlines()
    items = []
    for i, line in enumerate(lines):
        m = _NS_RE.match(line)
        if not m:
            continue
        nxt = lines[i + 1].strip() if i + 1 < len(lines) else ""
        if not re.search(r"\b(box|pcs?)\b", nxt, re.I):
            continue  # must have a pack line to be this format
        name = m.group(2).strip()
        qty = _to_float(m.group(3))
        amount = _to_float(m.group(5))
        if not amount or not qty:
            continue
        is_pcs, per_box = _parse_pack_spec(nxt)
        divisor = qty if is_pcs else qty * per_box
        name = _strip_packaging(re.sub(r"\s{2,}", " ", name).strip(" ,-*"))
        if len(name) < 2:
            continue
        items.append({
            "name": name,
            "qty": divisor,
            "cost": amount,
            "pricing_type": "quantity",
            "raw": (line + " | " + nxt).strip(),
        })
    return items


def _parse_text_lines(text):
    """Fallback for free-text (PDF/Word with no usable tables): one line each."""
    items = []
    for line in (text or "").splitlines():
        line = line.strip()
        if len(line) < 3 or _looks_like_header_or_total(line):
            continue
        money = _money_values(line)
        if not money:
            continue
        # name = the line with the money tokens stripped from the end
        name = MONEY_RE.sub("", line).strip(" .-\t")
        name = re.sub(r"\s{2,}", " ", name)
        if len(name) < 2 or name.replace(".", "").isdigit():
            continue
        cost = max(money)
        qty = None
        # a small leading integer is often a quantity: "12 Roma Tomato 24.00"
        lead = re.match(r"^(\d{1,4})\s+(.*)", name)
        if lead:
            qty = _to_float(lead.group(1))
            name = lead.group(2).strip()
        items.append({"name": name, "qty": qty, "cost": cost, "raw": line})
    return items


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------

def _extract_pdf(path):
    import pdfplumber
    text_chunks = []
    table_items = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_chunks.append(page.extract_text() or "")
            for table in (page.extract_tables() or []):
                table_items.extend(_parse_table_rows(table))
    full_text = "\n".join(text_chunks)

    # 1) Structured distributor-invoice rows are most reliable when present.
    #    Try the known per-line formats and use whichever recognizes more rows.
    structured = max(
        (_parse_structured_invoice(full_text),
         _parse_ziyad_invoice(full_text),
         _parse_karabetian(full_text),
         _parse_natures_sweets(full_text)),
        key=len,
    )
    if len(structured) >= 2:
        return structured
    # 2) Otherwise real tables, if pdfplumber found usable ones.
    if len(table_items) >= 2:
        return table_items
    # 3) Last resort: line-by-line text heuristic.
    return _parse_text_lines(full_text)


def _extract_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True, read_only=True)
    rows = []
    for ws in wb.worksheets:
        for r in ws.iter_rows(values_only=True):
            rows.append(list(r))
    return _parse_table_rows(rows)


def _extract_csv(path):
    rows = []
    with open(path, newline="", encoding="utf-8-sig", errors="replace") as f:
        for r in csv.reader(f):
            rows.append(r)
    return _parse_table_rows(rows)


def _extract_docx(path):
    from docx import Document
    doc = Document(path)
    items = []
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        items.extend(_parse_table_rows(rows))
    if len(items) < 2:
        text = "\n".join(p.text for p in doc.paragraphs)
        items.extend(_parse_text_lines(text))
    return items


def _dedupe(items):
    """Merge obvious duplicate names (sum qty, keep first cost)."""
    seen = {}
    order = []
    for it in items:
        key = " ".join(it["name"].lower().split())
        if key in seen:
            continue  # keep first occurrence; review screen lets user adjust
        seen[key] = it
        order.append(key)
    return [seen[k] for k in order]


def extract_invoice(path):
    """Dispatch on file extension and return a list of candidate line items."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        items = _extract_pdf(path)
    elif ext in (".xlsx", ".xlsm"):
        items = _extract_xlsx(path)
    elif ext in (".csv", ".tsv"):
        items = _extract_csv(path)
    elif ext in (".docx",):
        items = _extract_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    # clean up names
    for it in items:
        it["name"] = re.sub(r"\s{2,}", " ", (it["name"] or "").strip())
    return _dedupe([it for it in items if it["name"]])


SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".xlsm", ".csv", ".tsv", ".docx"}
